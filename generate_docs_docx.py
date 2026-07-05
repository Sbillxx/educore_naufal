import pymysql
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_docx():
    # Setup DB connection
    connection = pymysql.connect(
        host='127.0.0.1',
        user='root',
        password='',
        database='sistem_akademik',
        port=3306,
        cursorclass=pymysql.cursors.DictCursor
    )

    try:
        with connection.cursor() as cursor:
            sql = """
                SELECT u.role, u.name, u.email, 
                       COALESCE(c_stu.name, c_teach.name, '-') as class_name
                FROM users u
                LEFT JOIN students s ON u.id = s.user_id
                LEFT JOIN classes c_stu ON s.class_id = c_stu.id
                LEFT JOIN teachers t ON u.id = t.user_id
                LEFT JOIN classes c_teach ON t.id = c_teach.homeroom_teacher_id
                ORDER BY u.role, class_name, u.name
            """
            cursor.execute(sql)
            users = cursor.fetchall()
            
            # Create Document
            doc = Document()
            
            # Title
            title = doc.add_heading('Dokumentasi Akun Pengguna (Username & Password)', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph(
                "Berikut adalah daftar seluruh akun yang dapat digunakan untuk login ke dalam sistem, "
                "mencakup Guru Mapel, Wali Kelas, dan Siswa Riil."
            )
            
            p = doc.add_paragraph()
            run = p.add_run("Catatan: ")
            run.bold = True
            p.add_run("Seluruh akun di bawah ini menggunakan password default: ")
            p.add_run("password").bold = True
            
            roles = {
                'wali_kelas': 'Wali Kelas',
                'guru_mapel': 'Guru Mata Pelajaran',
                'siswa': 'Siswa'
            }
            
            for role_key, role_name in roles.items():
                role_users = [u for u in users if u['role'] == role_key]
                
                doc.add_heading(role_name, level=2)
                
                # Create table
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                
                # Header row
                hdr_cells = table.rows[0].cells
                headers = ['No', 'Nama Lengkap', 'Kelas / Keterangan', 'Email (Username)', 'Password']
                for i, text in enumerate(headers):
                    hdr_cells[i].text = text
                    # Bold headers
                    for paragraph in hdr_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Populate rows
                for i, u in enumerate(role_users, 1):
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(i)
                    row_cells[1].text = u['name']
                    row_cells[2].text = u['class_name']
                    row_cells[3].text = u['email']
                    row_cells[4].text = "password"
            
            output_path = os.path.join(r"C:\Users\ibnus\OneDrive\Documents\vscode\js\educore_naufal", "Dokumentasi_Akun_Pengguna_v2.docx")
            doc.save(output_path)
            print(f"Docx successfully saved to {output_path}")

    finally:
        connection.close()

if __name__ == "__main__":
    generate_docx()
